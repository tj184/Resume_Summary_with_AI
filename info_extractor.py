import os
from PyPDF2 import PdfReader
from docx import Document
from reportlab.pdfgen import canvas
from tempfile import NamedTemporaryFile
from llama import prompt

def extract_text_from_pdf(pdf_path):
    text = ""
    with open(pdf_path, 'rb') as file:
        reader = PdfReader(file)
        for page in reader.pages:
            text += page.extract_text()
    return text

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def convert_docx_to_pdf(docx_path):
    temp_pdf = NamedTemporaryFile(delete=False, suffix=".pdf")
    temp_pdf.close()

    c = canvas.Canvas(temp_pdf.name)
    doc = Document(docx_path)
    
    for paragraph in doc.paragraphs:
        c.drawString(100, 750, paragraph.text)
        c.showPage()
    c.save()

    return temp_pdf.name

def extract_text_from_files_in_folder(folder_path, role):
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        
        if filename.lower().endswith(".pdf"):
            print(f"Extracting text from PDF: {filename}")
            text = extract_text_from_pdf(file_path)
        
        elif filename.lower().endswith(".docx"):
            print(f"Converting and extracting text from DOCX: {filename}")
            pdf_path = convert_docx_to_pdf(file_path)
            text = extract_text_from_pdf(pdf_path)
            prompt(text, role)
        
        else:
            continue

folder_path = 'resume_pdf'
role = "Engineer"
extract_text_from_files_in_folder(folder_path, role)
